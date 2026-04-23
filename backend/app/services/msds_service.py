"""
化安通 (HuaAnTong) - MSDS自动生成系统
MSDS 文档服务层 - 封装 MSDS 生成、导出、审查功能
"""

import json
import logging
import threading
from datetime import datetime
from typing import Optional, Dict, Any, List

from sqlalchemy.orm import Session
from sqlalchemy import or_

from app.models.msds import MSDSDocument
from app.core.msds_pipeline import MSDSGenerator
from app.core.msds_reviewer import MSDSReviewer
from app.core.mixture_calculator import MixtureCalculator, Component, build_component
from app.config import settings

logger = logging.getLogger(__name__)


# ============================================================
# 后台任务存储（用于跟踪异步任务进度）
# ============================================================

_task_store: Dict[int, Dict[str, Any]] = {}


class MSDSService:
    """MSDS 文档服务"""

    def __init__(self, db: Session):
        self.db = db
        self.generator = MSDSGenerator()

    # ----------------------------------------------------------
    # 纯物质 MSDS 生成
    # ----------------------------------------------------------

    def generate_pure(self, cas_or_name: str, company_info: Optional[Dict] = None) -> int:
        """
        异步生成纯物质 MSDS

        Args:
            cas_or_name: CAS号或化学品名称
            company_info: 企业信息

        Returns:
            文档ID（任务ID）
        """
        # 创建数据库记录
        doc = MSDSDocument(
            title=f"MSDS - {cas_or_name}",
            cas_number=cas_or_name.split()[0] if cas_or_name else "",
            doc_type="pure",
            status="generating",
            company_info=json.dumps(company_info or {}, ensure_ascii=False),
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        task_id = doc.id
        _task_store[task_id] = {"status": "generating", "progress": "已创建任务"}

        # 启动后台线程
        thread = threading.Thread(
            target=self._generate_pure_task,
            args=(task_id, cas_or_name, company_info),
            daemon=True,
        )
        thread.start()

        return task_id

    def _generate_pure_task(self, task_id: int, cas_or_name: str, company_info: Optional[Dict]):
        """后台任务：生成纯物质MSDS"""
        try:
            _task_store[task_id] = {"status": "generating", "progress": "开始生成..."}

            def progress_callback(msg: str):
                _task_store[task_id]["progress"] = msg

            # 生成MSDS数据
            msds_data = self.generator.generate(cas_or_name, company_info, progress_callback)

            # 转为Markdown
            markdown_content = self.generator.to_markdown(msds_data)

            # 保存到文件
            cas = cas_or_name.strip().replace("-", "_").replace(" ", "_")
            output_dir = settings.output_dir / "pure"
            output_dir.mkdir(parents=True, exist_ok=True)
            output_file = output_dir / f"MSDS_{cas}_detailed.md"
            output_file.write_text(markdown_content, encoding="utf-8")

            # 更新数据库
            doc = self.db.query(MSDSDocument).filter(MSDSDocument.id == task_id).first()
            if doc:
                doc.status = "completed"
                doc.data_json = json.dumps(msds_data, ensure_ascii=False)
                doc.markdown_content = markdown_content
                doc.title = f"MSDS - {msds_data.get('part1_identification', {}).get('product_name_cn', cas_or_name)}"
                doc.cas_number = msds_data.get('part1_identification', {}).get('cas_number', '')
                self.db.commit()

            _task_store[task_id] = {
                "status": "completed",
                "progress": "生成完成",
                "output_file": str(output_file),
            }

            logger.info(f"MSDS生成完成: task_id={task_id}, file={output_file}")

        except Exception as e:
            logger.error(f"MSDS生成失败: task_id={task_id}, error={e}", exc_info=True)
            _task_store[task_id] = {"status": "failed", "progress": f"生成失败: {str(e)}"}

            doc = self.db.query(MSDSDocument).filter(MSDSDocument.id == task_id).first()
            if doc:
                doc.status = "failed"
                doc.error_message = str(e)
                self.db.commit()

    # ----------------------------------------------------------
    # 混合物 MSDS 生成
    # ----------------------------------------------------------

    def generate_mixture(
        self,
        product_name: str,
        components: List[Dict],
        company_info: Optional[Dict] = None,
    ) -> int:
        """
        异步生成混合物 MSDS

        Args:
            product_name: 产品名称
            components: 组分列表 [{"name": ..., "cas": ..., "concentration": ...}]
            company_info: 企业信息

        Returns:
            文档ID（任务ID）
        """
        doc = MSDSDocument(
            title=f"MSDS - {product_name}",
            doc_type="mixture",
            status="generating",
            company_info=json.dumps(company_info or {}, ensure_ascii=False),
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        task_id = doc.id
        _task_store[task_id] = {"status": "generating", "progress": "已创建任务"}

        thread = threading.Thread(
            target=self._generate_mixture_task,
            args=(task_id, product_name, components, company_info),
            daemon=True,
        )
        thread.start()

        return task_id

    def _generate_mixture_task(
        self,
        task_id: int,
        product_name: str,
        components: List[Dict],
        company_info: Optional[Dict],
    ):
        """后台任务：生成混合物MSDS"""
        try:
            _task_store[task_id] = {"status": "genering", "progress": "构建组分数据..."}

            # 构建组分对象
            comp_objects = []
            for comp in components:
                c = build_component(
                    name=comp.get("name", ""),
                    cas=comp.get("cas", ""),
                    concentration=comp.get("concentration", 0),
                )
                comp_objects.append(c)

            # 使用第一个主要组分作为CAS参考
            primary_cas = components[0].get("cas", "") if components else ""

            # 计算 GHS 分类
            _task_store[task_id]["progress"] = "计算混合物GHS分类..."
            calculator = MixtureCalculator(comp_objects)
            calc_result = calculator.calculate_all()

            # 使用LLM生成MSDS
            _task_store[task_id]["progress"] = "调用LLM生成MSDS..."

            def progress_callback(msg: str):
                _task_store[task_id]["progress"] = msg

            msds_data = self.generator.generate(primary_cas or product_name, company_info, progress_callback)

            # 将混合物计算结果合并到MSDS数据中
            msds_data["part3_composition"] = {
                "substance_type": "混合物",
                "components": [
                    {
                        "name": comp.get("name", ""),
                        "cas": comp.get("cas", ""),
                        "concentration": f"{comp.get('concentration', 0)}%",
                    }
                    for comp in components
                ],
                "ghs_calculation": {
                    "classifications": calc_result.classifications,
                    "h_codes": calc_result.h_codes,
                    "signal_word": calc_result.signal_word,
                },
            }

            # 更新混合物的危险分类
            if calc_result.classifications:
                msds_data["part2_hazard"]["ghs_classifications"] = [
                    c["hazard"] for c in calc_result.classifications
                ]
                msds_data["part2_hazard"]["signal_word"] = calc_result.signal_word
                msds_data["part2_hazard"]["hazard_codes"] = calc_result.h_codes

            # 转为Markdown
            markdown_content = self.generator.to_markdown(msds_data)

            # 保存文件
            safe_name = product_name.replace(" ", "_").replace("/", "_")
            output_dir = settings.output_dir / "mixture"
            output_dir.mkdir(parents=True, exist_ok=True)
            output_file = output_dir / f"MSDS_{safe_name}_detailed.md"
            output_file.write_text(markdown_content, encoding="utf-8")

            # 更新数据库
            doc = self.db.query(MSDSDocument).filter(MSDSDocument.id == task_id).first()
            if doc:
                doc.status = "completed"
                doc.data_json = json.dumps(msds_data, ensure_ascii=False)
                doc.markdown_content = markdown_content
                doc.title = f"MSDS - {product_name}"
                doc.cas_number = primary_cas
                self.db.commit()

            _task_store[task_id] = {
                "status": "completed",
                "progress": "生成完成",
                "output_file": str(output_file),
            }

        except Exception as e:
            logger.error(f"混合物MSDS生成失败: task_id={task_id}, error={e}", exc_info=True)
            _task_store[task_id] = {"status": "failed", "progress": f"生成失败: {str(e)}"}

            doc = self.db.query(MSDSDocument).filter(MSDSDocument.id == task_id).first()
            if doc:
                doc.status = "failed"
                doc.error_message = str(e)
                self.db.commit()

    # ----------------------------------------------------------
    # 任务状态查询
    # ----------------------------------------------------------

    def get_task_status(self, task_id: int) -> Optional[Dict]:
        """获取任务状态"""
        task_info = _task_store.get(task_id, {})
        doc = self.db.query(MSDSDocument).filter(MSDSDocument.id == task_id).first()

        if not doc:
            return None

        result = {
            "task_id": task_id,
            "status": doc.status,
            "progress": task_info.get("progress", ""),
        }

        if doc.status == "completed":
            result["result"] = {
                "id": doc.id,
                "title": doc.title,
                "cas_number": doc.cas_number,
                "doc_type": doc.doc_type,
                "status": doc.status,
                "data_json": doc.data_json,
                "markdown_content": doc.markdown_content,
                "company_info": doc.company_info,
                "error_message": doc.error_message,
                "created_at": str(doc.created_at) if doc.created_at else None,
                "updated_at": str(doc.updated_at) if doc.updated_at else None,
            }

        return result

    # ----------------------------------------------------------
    # 文档列表与详情
    # ----------------------------------------------------------

    def list_documents(
        self,
        page: int = 1,
        page_size: int = 20,
        doc_type: Optional[str] = None,
        status: Optional[str] = None,
    ) -> tuple:
        """
        获取MSDS文档列表

        Returns:
            (items, total) 元组
        """
        query = self.db.query(MSDSDocument)

        if doc_type:
            query = query.filter(MSDSDocument.doc_type == doc_type)
        if status:
            query = query.filter(MSDSDocument.status == status)

        total = query.count()
        items = (
            query.order_by(MSDSDocument.updated_at.desc())
            .offset((page - 1) * page_size)
            .limit(page_size)
            .all()
        )

        return items, total

    def get_document(self, doc_id: int) -> Optional[MSDSDocument]:
        """获取MSDS文档详情"""
        return self.db.query(MSDSDocument).filter(MSDSDocument.id == doc_id).first()

    def get_document_markdown(self, doc_id: int) -> Optional[str]:
        """获取MSDS文档Markdown内容"""
        doc = self.get_document(doc_id)
        if doc:
            return doc.markdown_content
        return None

    # ----------------------------------------------------------
    # 导出
    # ----------------------------------------------------------

    def export_pdf(self, doc_id: int) -> Optional[str]:
        """
        导出MSDS为PDF

        Returns:
            PDF文件路径
        """
        doc = self.get_document(doc_id)
        if not doc or not doc.markdown_content:
            return None

        try:
            import markdown as md_lib
            from weasyprint import HTML

            # Markdown -> HTML
            html_content = md_lib.markdown(
                doc.markdown_content,
                extensions=["tables", "fenced_code"],
            )

            # 添加样式
            styled_html = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    body {{ font-family: "Microsoft YaHei", "SimHei", sans-serif; margin: 40px; font-size: 14px; }}
    h1 {{ text-align: center; color: #333; }}
    h2 {{ color: #0066cc; border-bottom: 2px solid #0066cc; padding-bottom: 5px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
    th {{ background-color: #f5f5f5; }}
    hr {{ border: 1px solid #ccc; margin: 20px 0; }}
    strong {{ color: #333; }}
</style>
</head>
<body>
{html_content}
</body>
</html>
"""
            # HTML -> PDF
            output_dir = settings.output_dir / doc.doc_type
            output_dir.mkdir(parents=True, exist_ok=True)
            pdf_path = output_dir / f"MSDS_{doc.id}.pdf"
            HTML(string=styled_html).write_pdf(str(pdf_path))

            return str(pdf_path)

        except Exception as e:
            logger.error(f"PDF导出失败: {e}", exc_info=True)
            return None

    def export_word(self, doc_id: int) -> Optional[str]:
        """
        导出MSDS为Word

        Returns:
            Word文件路径
        """
        doc = self.get_document(doc_id)
        if not doc or not doc.markdown_content:
            return None

        try:
            from docx import Document
            from docx.shared import Pt, Inches

            document = Document()
            document.add_heading("化学品安全技术说明书（SDS）", level=0)

            # 按行处理Markdown
            lines = doc.markdown_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                if line.startswith("## "):
                    document.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    document.add_heading(line[2:], level=1)
                elif line.startswith("---"):
                    document.add_paragraph("")
                elif line.startswith("| ") and "|" in line[2:]:
                    # 简化处理表格
                    cells = [c.strip() for c in line.split("|")[1:-1]]
                    if cells and not all(set(c) <= {"-", " "} for c in cells):
                        p = document.add_paragraph(" | ".join(cells))
                        p.style.font.size = Pt(10)
                elif line.startswith("**") and "**" in line[2:]:
                    bold_end = line.index("**", 2)
                    bold_text = line[2:bold_end]
                    rest = line[bold_end + 2:].strip()
                    p = document.add_paragraph()
                    run = p.add_run(bold_text)
                    run.bold = True
                    if rest:
                        p.add_run(rest)
                elif line.startswith("- "):
                    document.add_paragraph(line[2:], style="List Bullet")
                else:
                    document.add_paragraph(line)

            output_dir = settings.output_dir / doc.doc_type
            output_dir.mkdir(parents=True, exist_ok=True)
            docx_path = output_dir / f"MSDS_{doc.id}.docx"
            document.save(str(docx_path))

            return str(docx_path)

        except Exception as e:
            logger.error(f"Word导出失败: {e}", exc_info=True)
            return None

    # ----------------------------------------------------------
    # 审查
    # ----------------------------------------------------------

    def review_document(self, doc_id: int) -> Optional[Dict]:
        """
        审查MSDS文档

        Returns:
            审查结果字典
        """
        doc = self.get_document(doc_id)
        if not doc:
            return None

        if doc.markdown_content:
            result = MSDSReviewer.review_from_markdown(doc.markdown_content)
        elif doc.data_json:
            try:
                msds_data = json.loads(doc.data_json)
                result = MSDSReviewer.review_from_data(msds_data)
            except json.JSONDecodeError:
                return {"status": "ERROR", "message": "数据格式错误"}
        else:
            return {"status": "ERROR", "message": "无审查内容"}

        # 保存审查结果
        doc.review_result = json.dumps(result, ensure_ascii=False)
        self.db.commit()

        return result

    # ----------------------------------------------------------
    # 删除
    # ----------------------------------------------------------

    def delete_document(self, doc_id: int) -> bool:
        """删除MSDS文档"""
        doc = self.get_document(doc_id)
        if not doc:
            return False

        self.db.delete(doc)
        self.db.commit()

        # 清理任务存储
        _task_store.pop(doc_id, None)

        return True
